#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理模块
支持多种格式文档的解析、文本提取和分块处理
"""

import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from loguru import logger

# 文档处理相关库
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    logger.warning("PyPDF2未安装，无法处理PDF文件")

try:
    from docx import Document
except ImportError:
    Document = None
    logger.warning("python-docx未安装，无法处理DOCX文件")

class DocumentProcessor:
    """
    文档处理器类，用于处理各种格式的文档
    """
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        """
        初始化文档处理器
        
        Args:
            chunk_size (int): 文本块大小
            chunk_overlap (int): 文本块重叠大小
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = ['.txt', '.pdf', '.docx', '.md']
        
        logger.info(f"初始化文档处理器 - 块大小: {chunk_size}, 重叠: {chunk_overlap}")
    
    def is_supported_format(self, file_path: str) -> bool:
        """
        检查文件格式是否支持
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            bool: 是否支持该格式
        """
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_formats
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """
        从TXT文件提取文本
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            str: 提取的文本
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            logger.debug(f"成功提取TXT文件文本: {len(content)}字符")
            return content
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    content = file.read()
                logger.debug(f"使用GBK编码成功提取TXT文件文本: {len(content)}字符")
                return content
            except Exception as e:
                logger.error(f"提取TXT文件失败: {e}")
                return ""
        except Exception as e:
            logger.error(f"提取TXT文件失败: {e}")
            return ""
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        从PDF文件提取文本
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            str: 提取的文本
        """
        if not PyPDF2:
            logger.error("PyPDF2未安装，无法处理PDF文件")
            return ""
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"提取PDF第{page_num+1}页失败: {e}")
                        continue
            
            logger.debug(f"成功提取PDF文件文本: {len(text)}字符")
            return text
        except Exception as e:
            logger.error(f"提取PDF文件失败: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        从DOCX文件提取文本
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            str: 提取的文本
        """
        if not Document:
            logger.error("python-docx未安装，无法处理DOCX文件")
            return ""
        
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.debug(f"成功提取DOCX文件文本: {len(text)}字符")
            return text
        except Exception as e:
            logger.error(f"提取DOCX文件失败: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        从文件提取文本（自动识别格式）
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            str: 提取的文本
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return ""
        
        if not self.is_supported_format(file_path):
            logger.error(f"不支持的文件格式: {file_path}")
            return ""
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.txt' or file_ext == '.md':
            return self.extract_text_from_txt(file_path)
        elif file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            logger.error(f"未知文件格式: {file_ext}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """
        清理文本内容
        
        Args:
            text (str): 原始文本
            
        Returns:
            str: 清理后的文本
        """
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符（保留中文、英文、数字、基本标点）
        text = re.sub(r'[^\u4e00-\u9fff\w\s.,!?;:()\[\]{}"\'-]', '', text)
        
        # 移除多余的换行符
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def split_text_into_chunks(self, text: str) -> List[Dict[str, Any]]:
        """
        将文本分割成块
        
        Args:
            text (str): 输入文本
            
        Returns:
            List[Dict[str, Any]]: 文本块列表
        """
        if not text:
            return []
        
        # 清理文本
        cleaned_text = self.clean_text(text)
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(cleaned_text):
            # 计算当前块的结束位置
            end = start + self.chunk_size
            
            # 如果不是最后一块，尝试在句号、问号、感叹号处分割
            if end < len(cleaned_text):
                # 寻找最近的句子结束符
                sentence_end = -1
                for i in range(end, max(start + self.chunk_size // 2, end - 100), -1):
                    if cleaned_text[i] in '.!?。！？':
                        sentence_end = i + 1
                        break
                
                if sentence_end > 0:
                    end = sentence_end
            
            # 提取当前块
            chunk_text = cleaned_text[start:end].strip()
            
            if chunk_text:
                chunk = {
                    'id': chunk_id,
                    'text': chunk_text,
                    'start_pos': start,
                    'end_pos': end,
                    'length': len(chunk_text)
                }
                chunks.append(chunk)
                chunk_id += 1
            
            # 计算下一块的开始位置（考虑重叠）
            start = max(start + 1, end - self.chunk_overlap)
        
        logger.info(f"文本分块完成: {len(chunks)}个块")
        return chunks
    
    def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理文档：提取文本并分块
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            List[Dict[str, Any]]: 处理后的文档块列表
        """
        logger.info(f"开始处理文档: {file_path}")
        
        # 提取文本
        text = self.extract_text_from_file(file_path)
        if not text:
            logger.error(f"无法从文件提取文本: {file_path}")
            return []
        
        # 分块
        chunks = self.split_text_into_chunks(text)
        
        # 添加文档信息
        file_name = Path(file_path).name
        for chunk in chunks:
            chunk['source_file'] = file_name
            chunk['source_path'] = file_path
        
        logger.info(f"文档处理完成: {file_name}, 共{len(chunks)}个块")
        return chunks
    
    def process_text(self, text: str, source_name: str = "text_input") -> List[Dict[str, Any]]:
        """
        处理纯文本：分块
        
        Args:
            text (str): 输入文本
            source_name (str): 来源名称
            
        Returns:
            List[Dict[str, Any]]: 处理后的文本块列表
        """
        logger.info(f"开始处理文本: {source_name}")
        
        if not text:
            logger.error("输入文本为空")
            return []
        
        # 分块
        chunks = self.split_text_into_chunks(text)
        
        # 添加来源信息
        for chunk in chunks:
            chunk['source_file'] = source_name
            chunk['source_path'] = source_name
        
        logger.info(f"文本处理完成: {source_name}, 共{len(chunks)}个块")
        return chunks


def test_document_processor():
    """
    测试文档处理器功能
    """
    processor = DocumentProcessor(chunk_size=200, chunk_overlap=20)
    
    # 测试文本处理
    test_text = "这是一个测试文本。" * 50
    chunks = processor.process_text(test_text, "测试文本")
    print(f"✅ 文本处理测试: 生成{len(chunks)}个块")
    
    # 显示第一个块的信息
    if chunks:
        first_chunk = chunks[0]
        print(f"第一个块: {first_chunk['text'][:50]}...")
        print(f"块长度: {first_chunk['length']}")


if __name__ == "__main__":
    test_document_processor()